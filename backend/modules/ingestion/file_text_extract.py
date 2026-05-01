"""Extract plain text from uploaded files by extension."""

from __future__ import annotations

import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def extract_text_from_local_file(path: Path, original_filename: str) -> str:
    """Pick extractor by extension; fall back to UTF-8 text."""
    name = original_filename or path.name
    suffix = Path(name).suffix.lower()

    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix == ".docx":
        return _extract_docx(path)

    raw = path.read_bytes()
    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError:
        return raw.decode("utf-8", errors="replace")


def _extract_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        logger.exception("pypdf not installed")
        return ""

    try:
        reader = PdfReader(str(path))
        parts: list[str] = []
        for page in reader.pages:
            t = page.extract_text()
            if t and t.strip():
                parts.append(t.strip())
        return "\n\n".join(parts)
    except Exception as exc:
        logger.warning("PDF extract failed for %s: %s", path, exc)
        return ""


def _extract_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError:
        logger.exception("python-docx not installed")
        return ""

    try:
        doc = Document(str(path))
        parts: list[str] = []
        for p in doc.paragraphs:
            if p.text and p.text.strip():
                parts.append(p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception as exc:
        logger.warning("DOCX extract failed for %s: %s", path, exc)
        return ""
